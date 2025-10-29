"""Export helpers: DOCX and PDF generation plus architecture diagrams."""
from typing import Dict
import os
# reportlab imported lazily in export_to_pdf


# Import heavy/optional deps lazily to avoid hard import errors when not installed


def generate_architecture_diagram(tech_text: str, out_dir: str) -> str:
    os.makedirs(out_dir, exist_ok=True)
    try:
        from graphviz import Digraph
    except Exception:
        # graphviz python package missing; create a simple DOT file fallback
        dot_path = os.path.join(out_dir, 'architecture.dot')
        with open(dot_path, 'w', encoding='utf-8') as f:
            f.write('# Graphviz not available; raw connections:\n')
            import re
            edges = re.findall(r"([A-Za-z0-9_ \-]+)\s*->\s*([A-Za-z0-9_ \-]+)", tech_text)
            for a, b in edges:
                f.write(f"{a} -> {b}\n")
        return dot_path
    g = Digraph(format='png')
    # Extract arrow edges
    import re
    edges = re.findall(r"([A-Za-z0-9_ \-]+)\s*->\s*([A-Za-z0-9_ \-]+)", tech_text)
    nodes = set()
    for a, b in edges:
        nodes.add(a.strip())
        nodes.add(b.strip())
        g.edge(a.strip(), b.strip())

    # If no edges, try to find a simple components list
    if not edges:
        for m in re.findall(r"^-\s*(.+)$", tech_text, flags=re.MULTILINE):
            n = m.strip()
            if n:
                nodes.add(n)
                g.node(n)

    out_path = os.path.join(out_dir, 'architecture.png')
    g.render(filename=os.path.join(out_dir, 'architecture'), cleanup=True)
    return out_path


def export_to_docx(artifacts: Dict[str, str], out_docx: str = 'outputs/final_deliverable.docx') -> str:
    try:
        from docx import Document
    except Exception:
        raise RuntimeError('python-docx is not installed; cannot create DOCX')

    doc = Document()
    doc.add_heading('Final Deliverables', level=1)

    for name, path in artifacts.items():
        doc.add_heading(name.capitalize(), level=2)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                txt = f.read()
            for line in txt.splitlines():
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(f'(Missing: {path})')

    os.makedirs(os.path.dirname(out_docx) or '.', exist_ok=True)
    doc.save(out_docx)
    return out_docx


def export_to_pdf(artifacts: Dict[str, str], out_pdf: str = 'outputs/final_deliverable.pdf') -> str:
    os.makedirs(os.path.dirname(out_pdf) or '.', exist_ok=True)
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except Exception:
        raise RuntimeError('reportlab is not installed; cannot generate PDF')

    c = canvas.Canvas(out_pdf, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont('Helvetica-Bold', 16)
    c.drawString(50, y, 'Final Deliverables')
    y -= 30

    for name, path in artifacts.items():
        c.setFont('Helvetica-Bold', 12)
        c.drawString(50, y, name.capitalize())
        y -= 20
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                txt = f.read()
            for line in txt.splitlines():
                if y < 60:
                    c.showPage()
                    y = height - 50
                c.setFont('Helvetica', 10)
                c.drawString(60, y, line[:120])
                y -= 14
        else:
            c.setFont('Helvetica', 10)
            c.drawString(60, y, f'(Missing: {path})')
            y -= 14

    c.save()
    return out_pdf
